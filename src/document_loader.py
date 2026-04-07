import os
import fitz  # pymupdf
import pytesseract
from PIL import Image
from docx import Document
import openpyxl
import pandas as pd
import io
 
# Si usa Windows, ajuste esta ruta a donde instaló Tesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
 
# =========================
# EXTRAER TEXTO DE IMAGEN (OCR)
# =========================
def extract_text_from_image_bytes(image_bytes):
    try:
        image = Image.open(io.BytesIO(image_bytes))
        text = pytesseract.image_to_string(image, lang="spa+eng")
        return text.strip()
    except Exception as e:
        print(f"Error en OCR: {e}")
        return ""
 
# =========================
# PDF (texto + imágenes con OCR)
# =========================
def load_pdf(file_path):
    text = ""
    doc = fitz.open(file_path)
    for page in doc:
        # Texto normal
        page_text = page.get_text()
        text += page_text
 
        # Si la página tiene poco texto, intentar OCR sobre la página entera
        if len(page_text.strip()) < 50:
            pix = page.get_pixmap(dpi=200)
            image_bytes = pix.tobytes("png")
            ocr_text = extract_text_from_image_bytes(image_bytes)
            text += ocr_text
 
        # Imágenes embebidas dentro del PDF
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            ocr_text = extract_text_from_image_bytes(image_bytes)
            if ocr_text:
                text += "\n" + ocr_text
 
    return text.strip()
 
# =========================
# WORD (.docx) — texto + imágenes con OCR
# =========================
def load_docx(file_path):
    doc = Document(file_path)
    text = ""
 
    # Texto de párrafos
    for para in doc.paragraphs:
        text += para.text + "\n"
 
    # Texto de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
        text += "\n"
 
    # Imágenes embebidas en el docx
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_bytes = rel.target_part.blob
            ocr_text = extract_text_from_image_bytes(image_bytes)
            if ocr_text:
                text += "\n" + ocr_text
 
    return text.strip()
 
# =========================
# EXCEL (.xlsx)
# =========================
def load_excel(file_path):
    text = ""
    wb = openpyxl.load_workbook(file_path, data_only=True)
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        text += f"\n[Hoja: {sheet}]\n"
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
            if row_text.strip():
                text += row_text + "\n"
    return text.strip()
 
# =========================
# CSV
# =========================
def load_csv(file_path):
    df = pd.read_csv(file_path)
    return df.to_string(index=False)
 
# =========================
# IMAGEN SUELTA (OCR)
# =========================
def load_image(file_path):
    with open(file_path, "rb") as f:
        image_bytes = f.read()
    return extract_text_from_image_bytes(image_bytes)
 
# =========================
# ROUTER PRINCIPAL
# =========================
def load_document(file_path):
    ext = os.path.splitext(file_path)[1].lower()
 
    loaders = {
        ".pdf":  load_pdf,
        ".docx": load_docx,
        ".doc":  load_docx,
        ".xlsx": load_excel,
        ".xls":  load_excel,
        ".csv":  load_csv,
        ".png":  load_image,
        ".jpg":  load_image,
        ".jpeg": load_image,
        ".tiff": load_image,
        ".bmp":  load_image,
        ".webp": load_image,
    }
 
    loader = loaders.get(ext)
    if not loader:
        raise ValueError(f"Formato no soportado: {ext}. Formatos aceptados: {', '.join(loaders.keys())}")
 
    return loader(file_path)
